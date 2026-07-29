"""Generate the Atlas review-edition DOCX from its authoritative Markdown.

The layout uses the compact_reference_guide document preset. The generator is
deliberately deterministic: Markdown owns the text, while this script owns Word
styles, numbering, tables, page geometry, and pagination safeguards.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "review" / "Atlas_Enterprise_Platform_Course_Review_Edition.md"
OUTPUT = ROOT / "Atlas_Enterprise_Platform_Course_Review_Edition.docx"

# compact_reference_guide token map
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
HEADER_DISTANCE = Inches(0.492)
# Keep the footer closer to the page edge than the body margin. Although the
# numeric distance is smaller, this increases clearance between body and footer.
FOOTER_DISTANCE = Inches(0.35)
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM_DXA = 80
CELL_MARGIN_START_END_DXA = 120

FONT = "Calibri"
NAVY = RGBColor(31, 78, 121)
BLUE = RGBColor(46, 116, 181)
TEAL = RGBColor(46, 95, 100)
MUTED = RGBColor(102, 102, 102)
INK = RGBColor(32, 32, 32)
TABLE_FILL = "E8EEF5"


def set_run_font(run, *, size=None, color=INK, bold=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for bullet_name in ("List Bullet", "List Bullet 2"):
        bullet_style = styles[bullet_name]
        bullet_style.font.name = FONT
        bullet_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        bullet_style.font.size = Pt(11)
        bullet_style.font.color.rgb = INK
        bullet_style.paragraph_format.left_indent = Inches(0.375)
        bullet_style.paragraph_format.first_line_indent = Inches(-0.1875)
        bullet_style.paragraph_format.space_before = Pt(0)
        bullet_style.paragraph_format.space_after = Pt(2)
        bullet_style.paragraph_format.line_spacing = 1.1

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
        "Heading 4": (11, NAVY, 8, 4),
        "Heading 5": (10.5, NAVY, 7, 3),
        "Heading 6": (10, NAVY, 6, 3),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = False
        # LibreOffice interprets a present keepNext element as enabled even
        # when OOXML stores w:val="0". Remove the element instead.
        style_p_pr = style._element.get_or_add_pPr()
        keep_next = style_p_pr.find(qn("w:keepNext"))
        if keep_next is not None:
            style_p_pr.remove(keep_next)
        style.paragraph_format.keep_together = True


def configure_section(section) -> None:
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.right_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE
    section.different_first_page_header_footer = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def configure_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    label = paragraph.add_run("Atlas Course Review Edition  •  ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(paragraph)
    for run in paragraph.runs[1:]:
        set_run_font(run, size=9, color=MUTED)


def create_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "252")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def apply_bullet(paragraph, bullet_index: int) -> None:
    paragraph.style = "List Bullet" if bullet_index % 2 == 0 else "List Bullet 2"


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline(paragraph, text: str, *, base_size=11, base_color=INK) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = part.startswith("*") and part.endswith("*") and not bold
        code = part.startswith("`") and part.endswith("`")
        clean = part[2:-2] if bold else part[1:-1] if italic or code else part
        run = paragraph.add_run(clean)
        set_run_font(
            run,
            size=base_size,
            color=base_color,
            bold=bold,
            italic=italic,
            name="Consolas" if code else FONT,
        )


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM_DXA),
        ("bottom", CELL_MARGIN_TOP_BOTTOM_DXA),
        ("start", CELL_MARGIN_START_END_DXA),
        ("end", CELL_MARGIN_START_END_DXA),
    ):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    tbl_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_width.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    normalized = [row + [""] * (columns - len(row)) for row in rows]
    if columns == 2:
        widths = [2700, 6660]
    elif columns == 3:
        widths = [1728, 4032, 3600]
    elif columns == 5:
        widths = [1872] * 5
    else:
        base = CONTENT_WIDTH_DXA // columns
        widths = [base] * columns
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)

    table = doc.add_table(rows=len(normalized), cols=columns)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for row_index, values in enumerate(normalized):
        if row_index == 0:
            header_repeat = OxmlElement("w:tblHeader")
            header_repeat.set(qn("w:val"), "true")
            table.rows[row_index]._tr.get_or_add_trPr().append(header_repeat)
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if row_index < len(normalized) - 1:
                paragraph.paragraph_format.keep_with_next = True
            add_inline(paragraph, value, base_size=10.5)
            if row_index == 0:
                shade_cell(cell, TABLE_FILL)
                for run in paragraph.runs:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Atlas Enterprise Platform")
    set_run_font(run, size=30, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Software Architecture Course")
    set_run_font(run, size=16, color=TEAL, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Review & Update Edition")
    set_run_font(run, size=22, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Condensed editorial control copy for structured review, maintenance, and future revision.")
    set_run_font(run, size=11, color=INK, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared July 29, 2026")
    set_run_font(run, size=10, color=MUTED)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_heading(doc: Document, level: int, text: str):
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 6)}")
    add_inline(
        paragraph,
        text,
        base_size={1: 16, 2: 13, 3: 12, 4: 11, 5: 10.5, 6: 10}.get(level, 10),
        base_color=BLUE if level < 3 else NAVY,
    )
    # Named pagination overrides for stable front-matter and teaching-brief
    # boundaries. Later chapters and appendices flow naturally so their Update
    # Notes do not become nearly empty continuation pages.
    if level == 1 and text.startswith("Chapter 1 "):
        paragraph.paragraph_format.page_break_before = True
    if level == 2 and text == "Course Map":
        paragraph.paragraph_format.page_break_before = True
    return paragraph


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section)
    configure_footer(section)
    add_cover(doc)
    current_chapter = ""
    current_part = ""
    bullet_index = 0

    # The first six Markdown lines are represented by the designed cover.
    index = next(i for i, line in enumerate(lines) if line.strip() == "## How to Use This Edition")
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line or line == "---":
            index += 1
            continue
        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            heading_text = heading.group(2)
            paragraph = add_heading(doc, level, heading_text)
            if level == 1:
                current_chapter = heading_text
                current_part = ""
            if level == 2:
                current_part = heading_text
            if current_chapter.startswith("Chapter 2 ") and heading_text == "Core Concepts":
                paragraph.paragraph_format.page_break_before = True
            if (
                current_chapter.startswith("Chapter 8 ")
                and current_part.startswith("Part 1 ")
                and heading_text == "Core Concepts"
            ):
                paragraph.paragraph_format.page_break_before = True
            index += 1
            continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.keep_together = True
            add_inline(paragraph, line[2:], base_size=11, base_color=NAVY)
            for run in paragraph.runs:
                run.bold = True
            index += 1
            continue
        if re.match(r"^-\s+", line):
            paragraph = doc.add_paragraph()
            apply_bullet(paragraph, bullet_index)
            bullet_index += 1
            add_inline(paragraph, re.sub(r"^-\s+", "", line))
            index += 1
            continue

        paragraph = doc.add_paragraph()
        add_inline(paragraph, line)
        index += 1

    doc.core_properties.title = "Atlas Enterprise Platform — Software Architecture Course Review & Update Edition"
    doc.core_properties.subject = "Generated editorial control copy"
    doc.core_properties.author = "Atlas Architecture Course"
    doc.core_properties.keywords = "Atlas, software architecture, review edition"
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT} from {SOURCE}")


if __name__ == "__main__":
    build()
